import os
import sys
import glob
from typing import Dict, Any, Optional, List
from langchain_community.vectorstores import FAISS
from langchain_core.embeddings import Embeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
import numpy as np
import importlib.util

LLM_MODEL_NAME = "gemma3:12b"
# LLM_MODEL_NAME = "llama3.1:8b"
EMBEDDING_MODEL_NAME = "nomic-embed-text"
VECTOR_STORE_PATH = "faiss_index"  # Path to save/load the FAISS index

class SimpleEmbeddings(Embeddings):
    """A simple embedding class that uses a basic approach to generate embeddings."""
    
    def __init__(self, dimensionality: int = 384):
        """Initialize with specified dimensionality."""
        self.dimensionality = dimensionality
    
    def _get_simple_embedding(self, text: str) -> List[float]:
        """Create a simple deterministic embedding based on the text content."""
        # Use a hash of the text to seed the random generator for consistency
        text_hash = hash(text) % (2**32)
        np.random.seed(text_hash)
        
        # Generate a random vector
        vector = np.random.normal(0, 1, self.dimensionality)
        
        # Normalize to unit length
        vector = vector / np.linalg.norm(vector)
        
        return vector.tolist()
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Embed a list of documents."""
        return [self._get_simple_embedding(text) for text in texts]
    
    def embed_query(self, text: str) -> List[float]:
        """Embed a query."""
        return self._get_simple_embedding(text)

def check_dependencies():
    """Check and inform about required dependencies for document processing."""
    missing_packages = []
    
    # Check for PDF support
    if importlib.util.find_spec("pypdf") is None:
        missing_packages.append("pypdf")
    
    # Check for DOCX support
    if importlib.util.find_spec("docx") is None:
        missing_packages.append("python-docx")
    
    if missing_packages:
        print("Missing dependencies for full document support:")
        print(f"To process PDF files: pip install {' '.join(missing_packages)}")
        print("You can still process Markdown files without these dependencies.")

def load_documents(resource_dir: str) -> List[Document]:
    """Load all supported document files from the resources directory."""
    # Define file patterns for different document types
    file_patterns = {
        "markdown": "*.md",
        "pdf": "*.pdf",
        "docx": "*.docx",
        "txt": "*.txt"
    }
    
    documents = []
    
    # Process each file type
    for doc_type, pattern in file_patterns.items():
        files = glob.glob(os.path.join(resource_dir, pattern))
        
        if not files:
            continue
            
        print(f"Found {len(files)} {doc_type} files")
        
        for file_path in files:
            try:
                file_name = os.path.basename(file_path)
                
                # Process based on file type
                if doc_type == "markdown" or doc_type == "txt":
                    with open(file_path, "r", encoding="utf-8") as file:
                        content = file.read()
                        doc = Document(page_content=content, metadata={"source": file_name})
                        documents.append(doc)
                
                elif doc_type == "pdf":
                    try:
                        from pypdf import PdfReader
                        reader = PdfReader(file_path)
                        content = ""
                        for page in reader.pages:
                            content += page.extract_text() + "\n"
                        doc = Document(page_content=content, metadata={"source": file_name})
                        documents.append(doc)
                    except ImportError:
                        print(f"Skipping {file_name}: pypdf module not installed. Install with: pip install pypdf")
                        continue
                
                elif doc_type == "docx":
                    try:
                        import docx
                        doc_reader = docx.Document(file_path)
                        content = "\n".join([para.text for para in doc_reader.paragraphs])
                        doc = Document(page_content=content, metadata={"source": file_name})
                        documents.append(doc)
                    except ImportError:
                        print(f"Skipping {file_name}: python-docx module not installed. Install with: pip install python-docx")
                        continue
                
                print(f"Loaded {file_name}")
            
            except Exception as e:
                print(f"Error loading {file_path}: {e}")
    
    return documents

def create_vector_store(documents: List[Document], force_recreate: bool = False) -> FAISS:
    """
    Create or load a vector store from documents.
    
    Args:
        documents: List of documents to process
        force_recreate: If True, recreate the vector store even if it exists
        
    Returns:
        FAISS vector store
    """
    base_folder = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    vector_store_path = os.path.join(base_folder, VECTOR_STORE_PATH)
    
    # If vector store exists and we're not forcing recreation, load it
    if os.path.exists(vector_store_path) and not force_recreate:
        try:
            print(f"Loading existing vector store from {vector_store_path}...")
            
            # Try to load embeddings
            try:
                from langchain_ollama import OllamaEmbeddings
            except ImportError:
                from langchain_community.embeddings import OllamaEmbeddings
                
            try:
                embeddings = OllamaEmbeddings(model=EMBEDDING_MODEL_NAME)
                # Test the embeddings
                _ = embeddings.embed_query("test")
            except Exception:
                print("Using simple local embeddings for loading vector store")
                embeddings = SimpleEmbeddings()
            
            # Load the vector store
            vector_store = FAISS.load_local(vector_store_path, embeddings)
            print(f"Successfully loaded vector store with {vector_store.index.ntotal} vectors")
            return vector_store
        except Exception as e:
            print(f"Error loading vector store: {e}")
            print("Creating new vector store...")
    
    # If we get here, we need to create a new vector store
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    
    chunks = text_splitter.split_documents(documents)
    print(f"Split documents into {len(chunks)} chunks")
    
    try:
        print(f"Trying to use Ollama embeddings with {EMBEDDING_MODEL_NAME}...")
        try:
            from langchain_ollama import OllamaEmbeddings
        except ImportError:
            from langchain_community.embeddings import OllamaEmbeddings
            
        embeddings = OllamaEmbeddings(model=EMBEDDING_MODEL_NAME)
        _ = embeddings.embed_query("test")
        print("Successfully using Ollama for embeddings")
    except Exception as e:
        print(f"Could not use Ollama embeddings: {e}")
        print("Falling back to simple local embeddings")
        embeddings = SimpleEmbeddings()
    
    vector_store = FAISS.from_documents(chunks, embeddings)
    
    # Save the vector store
    print(f"Saving vector store to {vector_store_path}...")
    vector_store.save_local(vector_store_path)
    print("Vector store saved successfully")
    
    return vector_store

def create_rag_chain(vector_store: FAISS) -> Dict[str, Any]:
    """
    Create a RAG chain using the vector store
    """
    # Create retriever
    retriever = vector_store.as_retriever(
        search_type="similarity",
        search_kwargs={"k": 5}  # Increased from 3 to 5 to get more context from diverse documents
    )
    
    # Create LLM using local Ollama with Gemma model
    # Import from newer package if available
    try:
        from langchain_ollama import OllamaLLM
        llm = OllamaLLM(model=LLM_MODEL_NAME)
    except ImportError:
        from langchain_community.llms import Ollama
        llm = Ollama(model=LLM_MODEL_NAME)
    
    # Create prompt template
    template = """
    You are an assistant that answers questions based on the provided context from multiple documents.
    
    Context:
    {context}
    
    Question: {question}
    
    Answer the question based only on the provided context. If you cannot answer the question from the context, say "I don't have enough information to answer this question."
    
    If relevant, mention which document(s) in the context contain information used for your answer.
    """
    
    prompt = ChatPromptTemplate.from_template(template)
    
    # Create RAG chain
    rag_chain = (
        {"context": retriever, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )
    
    return rag_chain

def main() -> None:
    import argparse
    
    parser = argparse.ArgumentParser(description="RAG system for document question answering")
    parser.add_argument("--rebuild", action="store_true", help="Force rebuild of the vector store")
    parser.add_argument("--query", type=str, help="Query to run against the RAG system")
    args = parser.parse_args()
    
    # Check for required dependencies
    check_dependencies()
    
    base_folder = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    resources_dir = os.path.join(base_folder, "resources")
    
    print(f"Loading documents from {resources_dir}...")
    documents = load_documents(resources_dir)
    
    if not documents:
        print("No documents found. Exiting.")
        return
    
    print(f"Creating/loading vector store with embeddings...")
    vector_store = create_vector_store(documents, force_recreate=args.rebuild)
    
    print(f"Creating RAG chain with Ollama LLM using {LLM_MODEL_NAME}...")
    rag_chain = create_rag_chain(vector_store)
    
    print(f"RAG system initialized with {len(documents)} documents")
    
    # Use provided query or default query
    query = args.query if args.query else "Describe the different applications of bioprinting technology across various fields"
    
    print(f"\nQuery: {query}")
    response = rag_chain.invoke(query)
    
    print("\nResponse:")
    print(response)

if __name__ == "__main__":
    main()